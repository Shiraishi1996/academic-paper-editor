import os
import json
import time
import re
import zipfile
import html as html_lib
import base64
import shutil
import tempfile
import traceback
import threading
import uuid
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from flask import Flask, render_template, request, jsonify, Response, send_file, send_from_directory, g, has_request_context
from werkzeug.exceptions import HTTPException
from werkzeug.utils import secure_filename
import docx
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
import openpyxl
from markdownify import markdownify as md
from lxml import etree, html as lxml_html

try:
    from deep_translator import GoogleTranslator
except Exception:
    GoogleTranslator = None

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = int(os.environ.get("MAX_UPLOAD_MB", "200")) * 1024 * 1024

# --- Configurations ---
BASE_DIR = os.path.abspath(os.path.dirname(__file__))


def _is_writable_dir(path):
    """Create/check a writable runtime directory without assuming the app folder is writable."""
    try:
        os.makedirs(path, exist_ok=True)
        probe = os.path.join(path, f".write_test_{os.getpid()}_{uuid.uuid4().hex}")
        with open(probe, "w", encoding="utf-8") as f:
            f.write("ok")
        os.remove(probe)
        return True
    except Exception:
        return False


def _resolve_runtime_root():
    """Prefer APP_DATA_DIR; fall back to app dir, then /tmp for serverless/read-only hosts."""
    candidates = []
    env_root = (os.environ.get("APP_DATA_DIR") or "").strip()
    if env_root:
        candidates.append(os.path.abspath(env_root))
    candidates.append(BASE_DIR)
    candidates.append(os.path.join(tempfile.gettempdir(), "academic_paper_editor"))
    for root in candidates:
        if _is_writable_dir(root):
            return root
    raise RuntimeError("No writable runtime directory is available. Set APP_DATA_DIR to a writable path.")


RUNTIME_ROOT = _resolve_runtime_root()
WORKSPACES_ROOT = os.path.join(RUNTIME_ROOT, "workspaces")
os.makedirs(WORKSPACES_ROOT, exist_ok=True)
_FILE_LOCK = threading.RLock()


def _workspace_id():
    if has_request_context():
        return getattr(g, "workspace_id", "default")
    return "default"


def _workspace_root():
    root = os.path.join(WORKSPACES_ROOT, _workspace_id())
    os.makedirs(root, exist_ok=True)
    return root


def _data_file():
    path = os.path.join(_workspace_root(), "data", "current_paper.json")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return path


def _snapshot_dir():
    path = os.path.join(_workspace_root(), "snapshots")
    os.makedirs(path, exist_ok=True)
    return path


def _upload_dir():
    path = os.path.join(_workspace_root(), "uploads")
    os.makedirs(path, exist_ok=True)
    return path


def _assets_json():
    return os.path.join(_upload_dir(), "assets_meta.json")


@app.before_request
def _bind_workspace():
    raw = (request.cookies.get("paper_editor_ws") or "").strip()
    if not re.fullmatch(r"[A-Za-z0-9_-]{16,64}", raw):
        raw = uuid.uuid4().hex
    g.workspace_id = raw


@app.after_request
def _persist_workspace_cookie(response):
    wid = getattr(g, "workspace_id", None)
    if wid and request.cookies.get("paper_editor_ws") != wid:
        response.set_cookie(
            "paper_editor_ws", wid, max_age=60 * 60 * 24 * 365,
            httponly=True, samesite="Lax", secure=request.is_secure
        )
    return response

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tif", ".tiff", ".svg"}


def new_id(prefix="id"):
    return f"{prefix}_{int(time.time() * 1000)}_{os.urandom(3).hex()}"


def make_block(block_type="paragraph", html="", plain_text="", source_index=None):
    return {
        "id": new_id("blk"),
        "type": block_type,
        "html": html,
        "plain_text": plain_text,
        "source_index": source_index,
    }


DEFAULT_PAPER = {
    "title": "",
    "authors": "",
    "affiliations": "",
    "abstract": "",
    "bib_data": "",
    "format_preset": "generic",
    "translation_provider": "online",
    "translation_model": "",
    "translation_glossary": "",
    "sections": [
        {
            "id": "sec_init",
            "title": "",
            "blocks": [make_block("paragraph", "<p></p>", "")],
            "content": "<p></p>",
        }
    ],
}


# --- HTML / structure helpers ---
def split_html_into_blocks(raw_html):
    """Convert legacy section HTML to individually editable top-level blocks."""
    if not raw_html or not raw_html.strip():
        return [make_block("paragraph", "<p></p>", "")]
    try:
        wrapper = lxml_html.fragment_fromstring(f"<div>{raw_html}</div>", create_parent=False)
        blocks = []
        if wrapper.text and wrapper.text.strip():
            t = wrapper.text.strip()
            blocks.append(make_block("paragraph", f"<p>{html_lib.escape(t)}</p>", t))
        for child in wrapper:
            fragment = lxml_html.tostring(child, encoding="unicode", method="html")
            tag = (child.tag or "").lower() if isinstance(child.tag, str) else ""
            if tag == "img" or (tag in {"p", "div", "figure"} and child.xpath(".//img") and not child.text_content().strip()):
                block_type = "image"
            elif tag == "table":
                block_type = "table"
            else:
                block_type = "paragraph"
            blocks.append(make_block(block_type, fragment, child.text_content().strip()))
            if child.tail and child.tail.strip():
                t = child.tail.strip()
                blocks.append(make_block("paragraph", f"<p>{html_lib.escape(t)}</p>", t))
        return blocks or [make_block("paragraph", raw_html, re.sub(r"<[^>]+>", "", raw_html))]
    except Exception:
        plain = re.sub(r"<[^>]+>", "", raw_html)
        return [make_block("paragraph", raw_html, plain)]


def sync_section_content(section):
    blocks = section.get("blocks") or []
    section["content"] = "\n".join((b.get("html") or "") for b in blocks)
    return section


def normalize_paper_structure(data):
    if not isinstance(data, dict):
        data = json.loads(json.dumps(DEFAULT_PAPER))
    data.setdefault("title", "")
    data.setdefault("authors", "")
    data.setdefault("affiliations", "")
    data.setdefault("abstract", "")
    data.setdefault("bib_data", "")
    data.setdefault("format_preset", "generic")
    data.setdefault("translation_provider", "online")
    data.setdefault("translation_model", "")
    data.setdefault("translation_glossary", "")
    # v18 uses one high-quality online translator only. Older local-provider projects are upgraded automatically.
    data["translation_provider"] = "online"
    data["translation_model"] = ""
    sections = data.setdefault("sections", [])
    if not sections:
        sections.append({"id": "sec_init", "title": "", "blocks": [make_block()], "content": "<p></p>"})
    for sec in sections:
        sec.setdefault("id", new_id("sec"))
        sec.setdefault("title", "")
        if not isinstance(sec.get("blocks"), list) or not sec.get("blocks"):
            sec["blocks"] = split_html_into_blocks(sec.get("content", ""))
        for blk in sec["blocks"]:
            blk.setdefault("id", new_id("blk"))
            blk.setdefault("type", "paragraph")
            blk.setdefault("html", "")
            blk.setdefault("plain_text", re.sub(r"<[^>]+>", "", blk.get("html", "")))
        sync_section_content(sec)
    return data


# --- Japanese academic style enforcement ---
# The editor is intended for manuscripts, so Japanese output is normalized to a plain academic style
# (常体・である調). This is deliberately conservative and targets sentence-final polite forms only.
_ACADEMIC_FIXED_REPLACEMENTS = [
    (r"ではありませんでした([。！？!?])", r"ではなかった\1"),
    (r"ではありません([。！？!?])", r"ではない\1"),
    (r"ではないです([。！？!?])", r"ではない\1"),
    (r"でございます([。！？!?])", r"である\1"),
    (r"でした([。！？!?])", r"であった\1"),
    (r"です([。！？!?])", r"である\1"),
    (r"ができます([。！？!?])", r"ができる\1"),
    (r"できます([。！？!?])", r"できる\1"),
    (r"必要があります([。！？!?])", r"必要がある\1"),
    (r"があります([。！？!?])", r"がある\1"),
    (r"ありません([。！？!?])", r"ない\1"),
    (r"なります([。！？!?])", r"なる\1"),
    (r"となります([。！？!?])", r"となる\1"),
    (r"と考えられます([。！？!?])", r"と考えられる\1"),
    (r"と推察されます([。！？!?])", r"と推察される\1"),
    (r"と示唆されます([。！？!?])", r"と示唆される\1"),
    (r"と判断されます([。！？!?])", r"と判断される\1"),
    (r"が示されます([。！？!?])", r"が示される\1"),
    (r"が認められます([。！？!?])", r"が認められる\1"),
    (r"が得られます([。！？!?])", r"が得られる\1"),
    (r"を示します([。！？!?])", r"を示す\1"),
    (r"を用います([。！？!?])", r"を用いる\1"),
    (r"を使用します([。！？!?])", r"を使用する\1"),
    (r"を比較します([。！？!?])", r"を比較する\1"),
    (r"を検討します([。！？!?])", r"を検討する\1"),
    (r"を評価します([。！？!?])", r"を評価する\1"),
    (r"を解析します([。！？!?])", r"を解析する\1"),
    (r"を実施します([。！？!?])", r"を実施する\1"),
    (r"を行います([。！？!?])", r"を行う\1"),
    (r"を提案します([。！？!?])", r"を提案する\1"),
    (r"を構築します([。！？!?])", r"を構築する\1"),
    (r"を算出します([。！？!?])", r"を算出する\1"),
    (r"を確認します([。！？!?])", r"を確認する\1"),
    (r"を明らかにします([。！？!?])", r"を明らかにする\1"),
    (r"しました([。！？!?])", r"した\1"),
    (r"されました([。！？!?])", r"された\1"),
    (r"いました([。！？!?])", r"いた\1"),
    (r"ありました([。！？!?])", r"あった\1"),
]


def to_japanese_academic_style(text):
    """Normalize Japanese sentence-final polite forms to manuscript-style plain form."""
    if not isinstance(text, str) or not text:
        return text or ""
    out = text
    for pattern, repl in _ACADEMIC_FIXED_REPLACEMENTS:
        out = re.sub(pattern, repl, out)
    # Common verbal polite endings that are safe enough to normalize at sentence boundaries.
    # Longer/specific rules above are applied first.
    common = {
        "示されています": "示されている", "報告されています": "報告されている",
        "考えています": "考えている", "用いています": "用いている",
        "含まれています": "含まれている", "知られています": "知られている",
    }
    for polite, plain in common.items():
        out = re.sub(re.escape(polite) + r"([。！？!?])", plain + r"\1", out)
    # Also normalize sentence-final polite forms when an editor fragment has no terminal punctuation.
    end_rules = [
        (r"ではありませんでした\s*$", "ではなかった"),
        (r"ではありません\s*$", "ではない"),
        (r"でした\s*$", "であった"),
        (r"です\s*$", "である"),
        (r"と考えられます\s*$", "と考えられる"),
        (r"と示唆されます\s*$", "と示唆される"),
        (r"されています\s*$", "されている"),
        (r"されました\s*$", "された"),
        (r"しました\s*$", "した"),
        (r"します\s*$", "する"),
        (r"できます\s*$", "できる"),
        (r"あります\s*$", "ある"),
        (r"ありません\s*$", "ない"),
        (r"なります\s*$", "なる"),
    ]
    for pattern, repl in end_rules:
        out = re.sub(pattern, repl, out)
    return out


def academic_style_html(raw_html):
    """Apply Japanese academic style to visible HTML text while preserving math/code/image markup."""
    if not isinstance(raw_html, str) or not raw_html.strip():
        return raw_html or ""
    try:
        root = lxml_html.fragment_fromstring(f"<div>{raw_html}</div>", create_parent=False)
        def walk(node, blocked=False):
            classes = set((node.get("class") or "").split()) if hasattr(node, "get") else set()
            tag = (node.tag or "").lower() if isinstance(node.tag, str) else ""
            here_blocked = blocked or tag in {"code", "pre", "script", "style"} or bool({"math-inline", "math-display", "equation-block"} & classes) or node.get("data-latex") is not None
            if not here_blocked and node.text:
                node.text = to_japanese_academic_style(node.text)
            for child in node:
                walk(child, here_blocked)
                if not here_blocked and child.tail:
                    child.tail = to_japanese_academic_style(child.tail)
        walk(root)
        return "".join(lxml_html.tostring(child, encoding="unicode", method="html") for child in root)
    except Exception:
        return to_japanese_academic_style(raw_html)


def enforce_academic_style(paper):
    """Enforce academic Japanese style throughout editable manuscript text, excluding references/authors/math/images."""
    result = normalize_paper_structure(deepcopy(paper or {}))
    result["abstract"] = to_japanese_academic_style(result.get("abstract", ""))
    ref_re = re.compile(r"^(references|bibliography|参考文献|引用文献)\s*$", re.I)
    for sec in result.get("sections", []):
        if ref_re.match((sec.get("title") or "").strip()):
            continue
        for blk in sec.get("blocks", []):
            if blk.get("type") in {"image", "equation"}:
                continue
            blk["html"] = academic_style_html(blk.get("html", ""))
            try:
                blk["plain_text"] = lxml_html.fromstring(f"<div>{blk['html']}</div>").text_content().strip()
            except Exception:
                blk["plain_text"] = re.sub(r"<[^>]+>", "", blk.get("html", ""))
        sync_section_content(sec)
    return result


# --- Utility Functions ---
def load_paper():
    with _FILE_LOCK:
        if os.path.exists(_data_file()):
            try:
                with open(_data_file(), "r", encoding="utf-8") as f:
                    return enforce_academic_style(json.load(f))
            except Exception as e:
                app.logger.warning("Paper data was unreadable and has been reset: %s", e)
                try:
                    broken = _data_file() + f".broken_{int(time.time())}"
                    shutil.copy2(_data_file(), broken)
                except Exception:
                    pass
        return enforce_academic_style(json.loads(json.dumps(DEFAULT_PAPER)))


def save_paper(data):
    normalized = enforce_academic_style(data)
    with _FILE_LOCK:
        os.makedirs(os.path.dirname(_data_file()), exist_ok=True)
        tmp = _data_file() + f".tmp_{os.getpid()}_{uuid.uuid4().hex}"
        try:
            with open(tmp, "w", encoding="utf-8") as f:
                json.dump(normalized, f, ensure_ascii=False, indent=2)
                f.flush()
                os.fsync(f.fileno())
            os.replace(tmp, _data_file())
        finally:
            if os.path.exists(tmp):
                try:
                    os.remove(tmp)
                except Exception:
                    pass
    return normalized


def load_assets():
    if os.path.exists(_assets_json()):
        try:
            with open(_assets_json(), "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return []


def save_assets(assets):
    with _FILE_LOCK:
        os.makedirs(os.path.dirname(_assets_json()), exist_ok=True)
        tmp = _assets_json() + f".tmp_{os.getpid()}_{uuid.uuid4().hex}"
        try:
            with open(tmp, "w", encoding="utf-8") as f:
                json.dump(assets, f, ensure_ascii=False, indent=2)
                f.flush()
                os.fsync(f.fileno())
            os.replace(tmp, _assets_json())
        finally:
            if os.path.exists(tmp):
                try:
                    os.remove(tmp)
                except Exception:
                    pass


def register_asset(assets, filename, original_name, extracted_text, is_image, source_document=None):
    asset = {
        "id": filename,
        "original_name": original_name,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "extracted_text": extracted_text,
        "is_image": is_image,
    }
    if source_document:
        asset["source_document"] = source_document
    # Avoid duplicate records when one embedded image is referenced repeatedly.
    assets[:] = [a for a in assets if a.get("id") != filename]
    assets.insert(0, asset)
    return asset


def extract_text(filepath, filename):
    ext = Path(filename).suffix.lower()
    text = ""
    try:
        if ext in [".txt", ".md", ".csv"]:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif ext == ".pdf":
            reader = PdfReader(filepath)
            for i, page in enumerate(reader.pages):
                extracted = page.extract_text()
                if extracted:
                    text += f"\n--- Page {i+1} ---\n" + extracted + "\n"
        elif ext == ".docx":
            doc = docx.Document(filepath)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif ext == ".xlsx":
            wb = openpyxl.load_workbook(filepath, data_only=True)
            for sheetname in wb.sheetnames:
                sheet = wb[sheetname]
                text += f"\n--- Sheet: {sheetname} ---\n"
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_data):
                        text += "\t".join(row_data) + "\n"
                text += "\n"
        else:
            text = "※テキスト抽出非対応のファイル形式です。"
    except Exception as e:
        text = f"[テキスト抽出エラー: {str(e)}]"
    if len(text) > 100000:
        text = text[:100000] + "\n... (省略)"
    return text


HEADING_REGEX = re.compile(
    r"^(?:"
    r"\d+(?:\.\d+)*[\.\-]?\s+.+|"
    r"[①-⑳]\s*.+|"
    r"[\(\[\{]\d+[\)\]\}]\s*.+|"
    r"第\d+[章節](?:\s+.*)?|"
    r"(?:概要|要旨|はじめに|序論|方法|結果|考察|結論|参考文献|Highlights|Abstract|Introduction|Method|Methods|Materials and Methods|Result|Results|Discussion|Conclusion|Conclusions|References|Author Contributions|Funding|Institutional Review Board Statement|Informed Consent Statement|Data Availability Statement|Conflicts of Interest|Abbreviations|Supplementary Materials)\s*[:：]?\s*"
    r")$",
    re.IGNORECASE,
)


def is_heading_paragraph(paragraph):
    text = (paragraph.text or "").strip()
    if not text:
        return False
    try:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name and (style_name.lower().startswith("heading") or "見出し" in style_name):
            return True
    except Exception:
        pass
    return bool(HEADING_REGEX.match(text))


def iter_block_items(parent):
    """Yield Paragraph and Table objects in original Word document order."""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        parent_part = parent.part
    else:
        parent_elm = parent._tc
        parent_part = parent.part
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def image_extension_from_part(image_part):
    ext = Path(str(getattr(image_part, "partname", ""))).suffix.lower()
    if ext in IMAGE_EXTS:
        return ext
    ctype = getattr(image_part, "content_type", "")
    return {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/webp": ".webp",
        "image/bmp": ".bmp",
        "image/tiff": ".tif",
        "image/svg+xml": ".svg",
    }.get(ctype, ".png")


def extract_embedded_image(part, rid, timestamp_prefix, image_cache, assets, source_document):
    cache_key = f"{getattr(part, 'partname', '')}:{rid}"
    if cache_key in image_cache:
        return image_cache[cache_key]
    try:
        image_part = part.related_parts[rid]
        ext = image_extension_from_part(image_part)
        img_filename = f"{timestamp_prefix}_img_{len(image_cache)+1:03d}{ext}"
        with open(os.path.join(_upload_dir(), img_filename), "wb") as f:
            f.write(image_part.blob)
        register_asset(
            assets,
            img_filename,
            f"{source_document} (画像 {len(image_cache)+1})",
            "[Word内画像]",
            True,
            source_document=source_document,
        )
        image_cache[cache_key] = img_filename
        return img_filename
    except Exception as e:
        print("Embedded image extraction error:", e)
        return None




# --- Publisher / manuscript style presets ---
PUBLISHER_PRESETS = {
    "generic": {
        "label": "Generic academic", "page": "A4", "font": "Times New Roman", "font_size": 11,
        "line_spacing": 1.15, "margins_cm": [2.54, 2.54, 2.54, 2.54], "columns": 1,
        "latex_class": "article", "latex_options": "a4paper,11pt"
    },
    "ieee": {
        "label": "IEEE", "page": "Letter", "font": "Times New Roman", "font_size": 10,
        "line_spacing": 1.0, "margins_cm": [1.9, 1.6, 1.9, 1.6], "columns": 2,
        "latex_class": "IEEEtran", "latex_options": "journal"
    },
    "elsevier": {
        "label": "Elsevier", "page": "A4", "font": "Times New Roman", "font_size": 11,
        "line_spacing": 1.15, "margins_cm": [2.5, 2.5, 2.5, 2.5], "columns": 1,
        "latex_class": "elsarticle", "latex_options": "preprint,12pt"
    },
    "springer_nature": {
        "label": "Springer Nature", "page": "A4", "font": "Times New Roman", "font_size": 10.5,
        "line_spacing": 1.15, "margins_cm": [2.5, 2.5, 2.5, 2.5], "columns": 1,
        "latex_class": "sn-jnl", "latex_options": "pdflatex,sn-basic"
    },
    "mdpi": {
        "label": "MDPI", "page": "A4", "font": "Times New Roman", "font_size": 10.5,
        "line_spacing": 1.15, "margins_cm": [2.0, 2.0, 2.0, 2.0], "columns": 1,
        "latex_class": "article", "latex_options": "a4paper,10pt"
    },
    "wiley": {
        "label": "Wiley", "page": "A4", "font": "Times New Roman", "font_size": 11,
        "line_spacing": 1.15, "margins_cm": [2.54, 2.54, 2.54, 2.54], "columns": 1,
        "latex_class": "article", "latex_options": "a4paper,11pt"
    },
    "taylor_francis": {
        "label": "Taylor & Francis", "page": "A4", "font": "Times New Roman", "font_size": 12,
        "line_spacing": 1.5, "margins_cm": [2.54, 2.54, 2.54, 2.54], "columns": 1,
        "latex_class": "article", "latex_options": "a4paper,12pt"
    },
}

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _local(el):
    try:
        return el.tag.split("}")[-1]
    except Exception:
        return ""

def _math_text(el):
    vals = []
    for node in el.iter():
        if _local(node) == 't' and (node.tag.startswith('{'+MATH_NS+'}') or node.tag == 't'):
            vals.append(node.text or '')
    return ''.join(vals)

def _latex_escape_math_text(text):
    """Convert common Unicode math characters emitted by Word OMML to portable LaTeX."""
    text = text or ''
    mapping = {
        'α': r'\alpha', 'β': r'\beta', 'γ': r'\gamma', 'δ': r'\delta', 'ε': r'\epsilon',
        'ζ': r'\zeta', 'η': r'\eta', 'θ': r'\theta', 'ι': r'\iota', 'κ': r'\kappa',
        'λ': r'\lambda', 'μ': r'\mu', 'ν': r'\nu', 'ξ': r'\xi', 'π': r'\pi', 'ρ': r'\rho',
        'σ': r'\sigma', 'τ': r'\tau', 'υ': r'\upsilon', 'φ': r'\phi', 'χ': r'\chi',
        'ψ': r'\psi', 'ω': r'\omega', 'Γ': r'\Gamma', 'Δ': r'\Delta', 'Θ': r'\Theta',
        'Λ': r'\Lambda', 'Ξ': r'\Xi', 'Π': r'\Pi', 'Σ': r'\Sigma', 'Φ': r'\Phi',
        'Ψ': r'\Psi', 'Ω': r'\Omega', '≤': r'\leq', '≥': r'\geq', '≠': r'\neq',
        '≈': r'\approx', '±': r'\pm', '∓': r'\mp', '×': r'\times', '·': r'\cdot',
        '∞': r'\infty', '∂': r'\partial', '∇': r'\nabla', '→': r'\to', '←': r'\leftarrow',
        '↔': r'\leftrightarrow', '∈': r'\in', '∉': r'\notin', '∝': r'\propto', '√': r'\sqrt{}',
        '\u2009': r'\,', '\u200a': r'\,', '\u00a0': '~',
    }
    out = []
    for ch in text:
        if ch in mapping:
            val = mapping[ch]
            if re.fullmatch(r'\\[A-Za-z]+', val):
                val += ' '
            out.append(val)
        elif ch == '\\':
            out.append(r'\backslash ')
        elif ch == '{':
            out.append(r'\{')
        elif ch == '}':
            out.append(r'\}')
        else:
            out.append(ch)
    return ''.join(out)

def omml_to_latex(el):
    """Best-effort OMML -> LaTeX converter for common Word equations. Exact OMML is preserved for DOCX re-export."""
    if el is None:
        return ''
    tag = _local(el)
    if tag in {'oMath', 'oMathPara', 'e', 'num', 'den', 'sub', 'sup', 'lim', 'deg'}:
        return ''.join(omml_to_latex(c) for c in el if _local(c) not in {'oMathParaPr', 'ctrlPr', 'rPr'})
    if tag == 'r':
        return _latex_escape_math_text(_math_text(el))
    if tag == 't':
        return _latex_escape_math_text(el.text or '')
    if tag == 'sSub':
        e = el.find(f'{{{MATH_NS}}}e'); sub = el.find(f'{{{MATH_NS}}}sub')
        return f"{{{omml_to_latex(e)}}}_{{{omml_to_latex(sub)}}}"
    if tag == 'sSup':
        e = el.find(f'{{{MATH_NS}}}e'); sup = el.find(f'{{{MATH_NS}}}sup')
        return f"{{{omml_to_latex(e)}}}^{{{omml_to_latex(sup)}}}"
    if tag == 'sSubSup':
        e = el.find(f'{{{MATH_NS}}}e'); sub = el.find(f'{{{MATH_NS}}}sub'); sup = el.find(f'{{{MATH_NS}}}sup')
        return f"{{{omml_to_latex(e)}}}_{{{omml_to_latex(sub)}}}^{{{omml_to_latex(sup)}}}"
    if tag == 'f':
        num = el.find(f'{{{MATH_NS}}}num'); den = el.find(f'{{{MATH_NS}}}den')
        return f"\\frac{{{omml_to_latex(num)}}}{{{omml_to_latex(den)}}}"
    if tag == 'rad':
        deg = el.find(f'{{{MATH_NS}}}deg'); e = el.find(f'{{{MATH_NS}}}e')
        d = omml_to_latex(deg)
        return f"\\sqrt[{d}]{{{omml_to_latex(e)}}}" if d else f"\\sqrt{{{omml_to_latex(e)}}}"
    if tag == 'd':
        pr = el.find(f'{{{MATH_NS}}}dPr')
        beg, end = '(', ')'
        if pr is not None:
            b = pr.find(f'{{{MATH_NS}}}begChr'); e2 = pr.find(f'{{{MATH_NS}}}endChr')
            if b is not None: beg = b.get(f'{{{MATH_NS}}}val', '(')
            if e2 is not None: end = e2.get(f'{{{MATH_NS}}}val', ')')
        body = ''.join(omml_to_latex(c) for c in el if _local(c) == 'e')
        return f"\\left{beg}{body}\\right{end}"
    if tag == 'acc':
        e = el.find(f'{{{MATH_NS}}}e'); pr = el.find(f'{{{MATH_NS}}}accPr')
        ch = ''
        if pr is not None:
            c = pr.find(f'{{{MATH_NS}}}chr')
            if c is not None: ch = c.get(f'{{{MATH_NS}}}val', '')
        amap = {'̇': 'dot', '̈': 'ddot', '^': 'hat', '̂': 'hat', '¯': 'bar', '̄': 'bar', '→': 'vec'}
        cmd = amap.get(ch)
        return f"\\{cmd}{{{omml_to_latex(e)}}}" if cmd else omml_to_latex(e)
    if tag == 'nary':
        pr = el.find(f'{{{MATH_NS}}}naryPr'); symbol = '∫'
        if pr is not None:
            c = pr.find(f'{{{MATH_NS}}}chr')
            if c is not None: symbol = c.get(f'{{{MATH_NS}}}val', symbol)
        cmd = {'∫': r'\int', '∑': r'\sum', '∏': r'\prod', '∮': r'\oint'}.get(symbol, symbol)
        sub = el.find(f'{{{MATH_NS}}}sub'); sup = el.find(f'{{{MATH_NS}}}sup'); e = el.find(f'{{{MATH_NS}}}e')
        out = cmd
        if sub is not None and omml_to_latex(sub): out += f"_{{{omml_to_latex(sub)}}}"
        if sup is not None and omml_to_latex(sup): out += f"^{{{omml_to_latex(sup)}}}"
        if e is not None: out += ' ' + omml_to_latex(e)
        return out
    if tag.endswith('Pr') or tag in {'ctrlPr', 'rPr', 'sty', 'chr', 'limLoc', 'begChr', 'endChr'}:
        return ''
    return ''.join(omml_to_latex(c) for c in el)

def math_span_from_omml(el, display=False):
    xml_bytes = etree.tostring(el, encoding='utf-8')
    b64 = base64.urlsafe_b64encode(xml_bytes).decode('ascii')
    latex = omml_to_latex(el).strip() or _math_text(el).strip()
    escaped_latex = html_lib.escape(latex, quote=True)
    body = f"\\[{escaped_latex}\\]" if display else f"\\({escaped_latex}\\)"
    cls = 'math-display' if display else 'math-inline'
    return f'<span class="{cls}" contenteditable="false" data-latex="{escaped_latex}" data-omml-b64="{b64}" data-math-display="{1 if display else 0}">{body}</span>', latex

def run_element_to_html(run_el, paragraph, timestamp_prefix, image_cache, assets, source_document):
    pieces = []
    text_nodes = [n.text or '' for n in run_el.iter(qn('w:t'))]
    if text_nodes:
        text = html_lib.escape(''.join(text_nodes)).replace('\n', '<br>')
        rpr = run_el.find(qn('w:rPr'))
        bold = rpr is not None and rpr.find(qn('w:b')) is not None
        italic = rpr is not None and rpr.find(qn('w:i')) is not None
        if bold: text = f'<strong>{text}</strong>'
        if italic: text = f'<em>{text}</em>'
        pieces.append(text)
    for blip in run_el.xpath('.//a:blip'):
        rid = blip.get(qn('r:embed')) or blip.get(qn('r:link'))
        if rid:
            fn = extract_embedded_image(paragraph.part, rid, timestamp_prefix, image_cache, assets, source_document)
            if fn:
                pieces.append(f'<img src="/uploads/{html_lib.escape(fn)}" alt="{html_lib.escape(source_document)} image" class="doc-image">')
    return ''.join(pieces)

def paragraph_to_html_and_assets(paragraph, timestamp_prefix, image_cache, assets, source_document):
    """Preserve paragraph order including Word OMML equations and embedded images."""
    pieces = []
    plain_parts = []
    images = []
    math_count = 0
    for child in paragraph._p:
        tag = _local(child)
        ns = child.tag.split('}')[0].strip('{') if '}' in child.tag else ''
        if ns == WORD_NS and tag == 'r':
            part_html = run_element_to_html(child, paragraph, timestamp_prefix, image_cache, assets, source_document)
            pieces.append(part_html)
            txt = ''.join((n.text or '') for n in child.iter(qn('w:t')))
            if txt: plain_parts.append(txt)
            for blip in child.xpath('.//a:blip'):
                rid = blip.get(qn('r:embed')) or blip.get(qn('r:link'))
                if rid:
                    fn = extract_embedded_image(paragraph.part, rid, timestamp_prefix, image_cache, assets, source_document)
                    if fn and fn not in images: images.append(fn)
        elif ns == MATH_NS and tag in {'oMath', 'oMathPara'}:
            span, latex = math_span_from_omml(child, display=(tag == 'oMathPara'))
            pieces.append(span); plain_parts.append(f'${latex}$'); math_count += 1
        elif ns == WORD_NS and tag == 'hyperlink':
            for r in child.findall(qn('w:r')):
                pieces.append(run_element_to_html(r, paragraph, timestamp_prefix, image_cache, assets, source_document))
                plain_parts.append(''.join((n.text or '') for n in r.iter(qn('w:t'))))

    seen = set(images)
    for blip in paragraph._p.xpath('.//a:blip'):
        rid = blip.get(qn('r:embed')) or blip.get(qn('r:link'))
        if rid:
            fn = extract_embedded_image(paragraph.part, rid, timestamp_prefix, image_cache, assets, source_document)
            if fn and fn not in seen:
                pieces.append(f'<img src="/uploads/{html_lib.escape(fn)}" alt="{html_lib.escape(source_document)} image" class="doc-image">')
                images.append(fn); seen.add(fn)

    text = ''.join(plain_parts).strip()
    if images and not text and not math_count:
        block_type = 'image'
        out_html = '<div class="figure-block">' + ''.join(pieces) + '</div>'
    elif math_count and not images:
        remainder = re.sub(r'\$[^$]*\$', '', text).strip()
        if not remainder or re.fullmatch(r'\(?\d+\)?', remainder):
            block_type = 'equation'
            out_html = '<div class="equation-block">' + ''.join(pieces) + '</div>'
        else:
            block_type = 'paragraph'
            out_html = '<p>' + ''.join(pieces) + '</p>'
    else:
        block_type = 'paragraph'
        out_html = '<p>' + ''.join(pieces) + '</p>'
    return block_type, out_html, text, images, math_count

def table_to_html(table, timestamp_prefix, image_cache, assets, source_document):
    rows_html = []
    table_math = 0
    equation_like = False
    for row in table.rows:
        cells_html = []
        row_has_math = 0
        for cell in row.cells:
            p_htmls = []
            cell_math = 0
            for p in cell.paragraphs:
                btype, ph, plain, _, mc = paragraph_to_html_and_assets(p, timestamp_prefix, image_cache, assets, source_document)
                cell_math += mc
                if ph and ph not in {'<p></p>', '<div class="equation-block"></div>'}:
                    p_htmls.append(ph)
            row_has_math += cell_math
            table_math += cell_math
            cells_html.append('<td>' + ''.join(p_htmls) + '</td>')
        if len(row.cells) in {2, 3} and row_has_math and any(re.match(r'^\s*\(?\d+\)?\s*$', (c.text or '').strip()) for c in row.cells):
            equation_like = True
        rows_html.append('<tr>' + ''.join(cells_html) + '</tr>')
    cls = 'doc-table equation-table' if equation_like else 'doc-table'
    return f'<table class="{cls}"><tbody>' + ''.join(rows_html) + '</tbody></table>', table_math


DOC_TYPE_LABELS = {"article", "review", "communication", "letter", "editorial", "case report", "perspective", "brief report"}
FRONT_STOP_LABELS = {"highlights", "abstract", "keywords", "introduction", "1. introduction"}


def _clean_front_text(text):
    return re.sub(r"\s+", " ", (text or "").strip())


def _looks_like_filename(text, original_name):
    t = _clean_front_text(text).lower()
    if not t:
        return False
    names = {original_name.lower(), Path(original_name).stem.lower()}
    return t in names


def extract_docx_front_matter(doc, original_name):
    """Best-effort metadata extraction that never uses the uploaded filename as the paper title."""
    nonempty = [(i, p, _clean_front_text(p.text)) for i, p in enumerate(doc.paragraphs) if _clean_front_text(p.text)]
    skip_elements = set()
    skip_indices = set()
    title = _clean_front_text(getattr(doc.core_properties, "title", ""))
    title_idx = None
    core_low = title.lower().rstrip(":") if title else ""
    if (not title or _looks_like_filename(title, original_name) or "type of the paper" in core_low or "type of paper" in core_low or core_low in DOC_TYPE_LABELS):
        title = ""

    # Prefer an explicit Word Title style when available.
    if not title:
        for i, p, text in nonempty[:30]:
            style = (p.style.name if p.style else "").lower()
            if "title" in style and text.lower() not in DOC_TYPE_LABELS and not _looks_like_filename(text, original_name):
                title, title_idx = text, i
                break

    # Template documents do not always use a Title style; use a conservative front-page heuristic.
    if not title:
        for i, p, text in nonempty[:30]:
            low = text.lower().rstrip(":")
            if low in DOC_TYPE_LABELS or low in {"type of the paper", "type of paper"}:
                continue
            if low.startswith(("academic editor", "received", "revised", "accepted", "published", "copyright")):
                continue
            if low in FRONT_STOP_LABELS or low.startswith("keywords"):
                break
            if _looks_like_filename(text, original_name) or "@" in text:
                continue
            # A manuscript title is usually longer than a person's name / affiliation line.
            if 10 <= len(text) <= 350 and (len(text.split()) >= 4 or re.search(r"[一-龥ぁ-んァ-ン]", text)):
                title, title_idx = text, i
                break

    if title and _looks_like_filename(title, original_name):
        title = ""
        title_idx = None

    authors = ""
    affiliations = ""
    abstract = ""

    # Locate author line immediately after the detected title.
    if title_idx is not None:
        candidates = [(i, p, t) for i, p, t in nonempty if i > title_idx][:8]
        for i, p, text in candidates:
            low = text.lower().rstrip(":")
            if low in FRONT_STOP_LABELS or low.startswith("keywords"):
                break
            if "@" in text or re.match(r"^[*\d]+\s*\t?", text):
                continue
            if len(text) <= 300:
                authors = text
                skip_elements.add(p._p)
                skip_indices.add(i)
                author_idx = i
                # affiliations / correspondence between author line and front-matter stop
                aff = []
                for j, q, qt in nonempty:
                    if j <= author_idx:
                        continue
                    qlow = qt.lower().rstrip(":")
                    if qlow in FRONT_STOP_LABELS or qlow.startswith("keywords"):
                        break
                    if j > author_idx + 8:
                        break
                    aff.append(qt)
                    skip_elements.add(q._p)
                    skip_indices.add(j)
                affiliations = "\n".join(aff)
                break

    # Extract Abstract body but keep Highlights/Keywords as manuscript content.
    abs_start = None
    for i, p, text in nonempty[:80]:
        if text.lower().rstrip(":") == "abstract":
            abs_start = i
            skip_elements.add(p._p)
            skip_indices.add(i)
            break
    if abs_start is not None:
        parts = []
        for i, p, text in nonempty:
            if i <= abs_start:
                continue
            low = text.lower()
            if low.startswith("keywords") or HEADING_REGEX.match(text):
                break
            parts.append(text)
            skip_elements.add(p._p)
            skip_indices.add(i)
        abstract = " ".join(parts).strip()

    # Remove generic document-type labels and the title paragraph from the editable body.
    for i, p, text in nonempty[:20]:
        low = text.lower().rstrip(":")
        if low in DOC_TYPE_LABELS or low in {"type of the paper", "type of paper"} or _looks_like_filename(text, original_name):
            skip_elements.add(p._p)
            skip_indices.add(i)
    if title_idx is not None:
        for i, p, text in nonempty:
            if i == title_idx:
                skip_elements.add(p._p)
                skip_indices.add(i)
                break

    return {"title": title, "authors": authors, "affiliations": affiliations, "abstract": abstract, "skip_elements": skip_elements, "skip_indices": skip_indices}



# --- Translation engine: higher-quality local Japanese/English models, online optional ---
# v17 defaults to FuguMT, a Japanese/English-specific model pair. M2M100 is available as a
# larger multilingual alternative. Both run locally after the first model download and require
# no API key. The older Helsinki OPUS-MT profile is kept only as a lightweight fallback.
LOCAL_MODEL_CACHE_DIR = os.path.abspath(os.environ.get("LOCAL_TRANSLATION_CACHE", os.path.join(RUNTIME_ROOT, "models")))
os.makedirs(LOCAL_MODEL_CACHE_DIR, exist_ok=True)
_LOCAL_TRANSLATION_LOCK = threading.RLock()
_LOCAL_TRANSLATION_MODELS = {}
LOCAL_TRANSLATION_PROFILES = {
    "fugumt": {
        "label": "FuguMT Japanese/English",
        "ja_en": os.environ.get("LOCAL_TRANSLATION_FUGU_JA_EN", "staka/fugumt-ja-en"),
        "en_ja": os.environ.get("LOCAL_TRANSLATION_FUGU_EN_JA", "staka/fugumt-en-ja"),
        "kind": "seq2seq",
        "beams": 8,
    },
    "m2m100": {
        "label": "M2M100 418M",
        "model": os.environ.get("LOCAL_TRANSLATION_M2M100", "facebook/m2m100_418M"),
        "kind": "m2m100",
        "beams": 6,
    },
    "opus": {
        "label": "OPUS-MT lightweight",
        "ja_en": os.environ.get("LOCAL_TRANSLATION_MODEL_JA_EN", "Helsinki-NLP/opus-mt-ja-en"),
        "en_ja": os.environ.get("LOCAL_TRANSLATION_MODEL_EN_JA", "Helsinki-NLP/opus-mt-en-jap"),
        "kind": "seq2seq",
        "beams": 5,
    },
}

# Generic academic terms. Users can add/override pairs in the editor with `日本語 = English`.
_BUILTIN_TRANSLATION_GLOSSARY = [
    ("提案手法", "proposed method"), ("従来手法", "conventional method"),
    ("実験結果", "experimental results"), ("数値シミュレーション", "numerical simulation"),
    ("統計的有意差", "statistical significance"), ("信頼区間", "confidence interval"),
    ("標準偏差", "standard deviation"), ("目的関数", "objective function"),
    ("有限要素法", "finite element method"), ("遺伝的アルゴリズム", "genetic algorithm"),
    ("地理情報システム", "geographic information system"), ("制御系", "control system"),
    ("深度制御", "depth control"), ("可変浮力", "variable buoyancy"),
    ("空気圧駆動", "pneumatic actuation"), ("時間遅延", "time delay"),
    ("オーバーシュート", "overshoot"), ("整定時間", "settling time"),
]


def parse_translation_glossary(glossary_text=""):
    pairs = {ja: en for ja, en in _BUILTIN_TRANSLATION_GLOSSARY}
    for raw in (glossary_text or "").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        parts = None
        for sep in ("=>", "->", "=", "\t"):
            if sep in line:
                parts = line.split(sep, 1); break
        if not parts:
            continue
        ja, en = (x.strip() for x in parts)
        if ja and en: pairs[ja] = en
    return list(pairs.items())


def _reference_target(text, target):
    if target == "en":
        text = re.sub(r"図\s*([（(]?\s*\d+(?:[-–—]\d+)?[a-zA-Z]?\s*[）)]?)", lambda m: f"Figure {m.group(1).strip('（）() ')}", text)
        text = re.sub(r"表\s*([（(]?\s*\d+(?:[-–—]\d+)?[a-zA-Z]?\s*[）)]?)", lambda m: f"Table {m.group(1).strip('（）() ')}", text)
        text = re.sub(r"式\s*[（(]?\s*(\d+(?:[-–—]\d+)?[a-zA-Z]?)\s*[）)]?", lambda m: f"Equation ({m.group(1)})", text)
    else:
        text = re.sub(r"\b(?:Fig(?:ure)?)[.]?\s*(\d+(?:[-–—]\d+)?[a-zA-Z]?)", lambda m: f"図{m.group(1)}", text, flags=re.I)
        text = re.sub(r"\bTable\s*(\d+(?:[-–—]\d+)?[a-zA-Z]?)", lambda m: f"表{m.group(1)}", text, flags=re.I)
        text = re.sub(r"\b(?:Eq(?:uation)?)[.]?\s*[（(]?\s*(\d+(?:[-–—]\d+)?[a-zA-Z]?)\s*[）)]?", lambda m: f"式（{m.group(1)}）", text, flags=re.I)
    return text


def _needs_translation(text, target):
    if not (text or "").strip(): return False
    has_ja = bool(re.search(r"[ぁ-んァ-ヶ一-龯々]", text))
    has_latin = bool(re.search(r"[A-Za-z]", text))
    return has_ja if target == "en" else has_latin


def _apply_glossary_target_terms(text, target, glossary_text=""):
    out = text or ""
    pairs = parse_translation_glossary(glossary_text)
    source_target = pairs if target == "en" else [(en, ja) for ja, en in pairs]
    for source, target_term in sorted(source_target, key=lambda x: len(x[0]), reverse=True):
        if not source: continue
        if re.search(r"[A-Za-z]", source) and not re.search(r"[ぁ-んァ-ヶ一-龯々]", source):
            rx = re.compile(r"(?i:(?<![A-Za-z0-9])" + re.escape(source) + r"(?![A-Za-z0-9]))")
        else:
            rx = re.compile(re.escape(source))
        out = rx.sub(lambda m, t=target_term: f"⟦{t}⟧", out)
    return out


def _protected_segments(text, target, glossary_text=""):
    text = _reference_target(text or "", target)
    protected_patterns = [
        r"https?://[^\s<>()]+", r"doi:\s*10\.\d{4,9}/[-._;()/:A-Za-z0-9]+",
        r"10\.\d{4,9}/[-._;()/:A-Za-z0-9]+", r"\[(?:\s*\d+\s*(?:[-–—,;]\s*\d+\s*)*)\]",
        r"(?i:\b(?:Figure|Fig\.?|Table|Equation|Eq\.?)\s*\(?\d+(?:[-–—]\d+)?[a-zA-Z]?\)?)",
        r"(?:図|表)\s*\d+(?:[-–—]\d+)?[a-zA-Z]?", r"式\s*[（(]?\s*\d+(?:[-–—]\d+)?[a-zA-Z]?\s*[）)]?",
        r"\b[A-Z][A-Z0-9-]{1,12}\b", r"\$[^$]+\$", r"\\\([^)]*\\\)", r"⟦[^⟧]+⟧",
    ]
    rx = re.compile("(" + "|".join(protected_patterns) + ")")
    out, pos = [], 0
    for m in rx.finditer(text):
        if m.start() > pos: out.append((True, text[pos:m.start()]))
        out.append((False, m.group(0))); pos = m.end()
    if pos < len(text): out.append((True, text[pos:]))
    return out


def _basic_sentence_chunks(value, target, max_chars=None):
    if not value: return []
    if max_chars is None: max_chars = 240 if target == "en" else 520
    if len(value) <= max_chars: return [value]
    pieces = re.split(r"(?<=[。！？!?])|(?<=[.])\s+(?=[A-Z0-9])|\n+", value)
    pieces = [p for p in pieces if p]
    chunks, buf = [], ""
    for p in pieces:
        if len(buf) + len(p) <= max_chars: buf += p; continue
        if buf: chunks.append(buf); buf = ""
        if len(p) <= max_chars: buf = p; continue
        rest = p
        while len(rest) > max_chars:
            cut = max(rest.rfind("、", 0, max_chars), rest.rfind(", ", 0, max_chars), rest.rfind("; ", 0, max_chars), rest.rfind(" ", 0, max_chars))
            cut = max_chars if cut < max_chars // 3 else cut + 1
            chunks.append(rest[:cut]); rest = rest[cut:]
        buf = rest
    if buf: chunks.append(buf)
    return chunks


def _local_translation_dependencies_available():
    try:
        import importlib.util
        return all(importlib.util.find_spec(m) is not None for m in ("transformers", "torch", "sentencepiece"))
    except Exception: return False


def _normalize_local_model_profile(profile):
    return profile if profile in LOCAL_TRANSLATION_PROFILES else "fugumt"


def _profile_model_ids(profile):
    p = LOCAL_TRANSLATION_PROFILES[_normalize_local_model_profile(profile)]
    if p["kind"] == "m2m100": return [p["model"]]
    return [p["ja_en"], p["en_ja"]]


def _local_model_cache_present(target="en", profile="fugumt"):
    p = LOCAL_TRANSLATION_PROFILES[_normalize_local_model_profile(profile)]
    model_id = p["model"] if p["kind"] == "m2m100" else (p["ja_en"] if target == "en" else p["en_ja"])
    marker = "models--" + model_id.replace("/", "--")
    return os.path.isdir(os.path.join(LOCAL_MODEL_CACHE_DIR, marker)) or os.path.isdir(os.path.join(LOCAL_MODEL_CACHE_DIR, model_id.replace("/", os.sep)))


def _load_local_translation_model(target, profile="fugumt"):
    if target not in {"ja", "en"}: raise ValueError("target must be ja or en")
    if not _local_translation_dependencies_available():
        raise RuntimeError("Local translation packages are not installed. Run: pip install -r requirements.txt")
    profile = _normalize_local_model_profile(profile)
    p = LOCAL_TRANSLATION_PROFILES[profile]
    model_id = p["model"] if p["kind"] == "m2m100" else (p["ja_en"] if target == "en" else p["en_ja"])
    key = (profile, model_id)
    with _LOCAL_TRANSLATION_LOCK:
        if key in _LOCAL_TRANSLATION_MODELS: return _LOCAL_TRANSLATION_MODELS[key]
        try:
            import torch
            from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
            offline = str(os.environ.get("LOCAL_TRANSLATION_OFFLINE", "0")).lower() in {"1","true","yes","on"}
            tokenizer = AutoTokenizer.from_pretrained(model_id, cache_dir=LOCAL_MODEL_CACHE_DIR, local_files_only=offline)
            model = AutoModelForSeq2SeqLM.from_pretrained(model_id, cache_dir=LOCAL_MODEL_CACHE_DIR, local_files_only=offline)
            device = "cuda" if torch.cuda.is_available() else "cpu"
            if device == "cuda": model = model.half()
            model.to(device); model.eval()
            _LOCAL_TRANSLATION_MODELS[key] = (tokenizer, model, device, p)
            return _LOCAL_TRANSLATION_MODELS[key]
        except Exception as e:
            if str(os.environ.get("LOCAL_TRANSLATION_OFFLINE", "0")).lower() in {"1","true","yes","on"}:
                raise RuntimeError(f"Local model is not cached for offline use: {model_id}") from e
            raise RuntimeError(f"Could not load/download local translation model {model_id}: {e}") from e


def _translate_local_piece(piece, target, profile="fugumt"):
    if not (piece or "").strip() or not _needs_translation(piece, target): return piece
    tokenizer, model, device, p = _load_local_translation_model(target, profile)
    import torch
    model_max = min(int(getattr(tokenizer, "model_max_length", 512) or 512), 1024)
    token_count = len(tokenizer.encode(piece, add_special_tokens=True))
    if token_count > int(model_max * .82) and len(piece) > 80:
        midpoint = len(piece)//2
        candidates=[piece.rfind("。",0,midpoint),piece.rfind(". ",0,midpoint),piece.rfind("、",0,midpoint),piece.rfind(" ",0,midpoint)]
        cut=max(candidates); cut=midpoint if cut < len(piece)//4 else cut+1
        return _translate_local_piece(piece[:cut],target,profile)+_translate_local_piece(piece[cut:],target,profile)
    if p["kind"] == "m2m100":
        tokenizer.src_lang = "ja" if target == "en" else "en"
        batch = tokenizer(piece, return_tensors="pt", truncation=True, max_length=model_max)
        batch = {k:v.to(device) for k,v in batch.items()}
        with torch.no_grad():
            generated = model.generate(**batch, forced_bos_token_id=tokenizer.get_lang_id(target), max_new_tokens=min(512, model_max), num_beams=p.get("beams",6), length_penalty=1.0, no_repeat_ngram_size=3, early_stopping=True)
    else:
        batch = tokenizer(piece, return_tensors="pt", truncation=True, max_length=model_max)
        batch = {k:v.to(device) for k,v in batch.items()}
        with torch.no_grad():
            generated = model.generate(**batch, max_new_tokens=min(512, model_max), num_beams=p.get("beams",8), length_penalty=1.0, no_repeat_ngram_size=3, early_stopping=True)
    return tokenizer.decode(generated[0], skip_special_tokens=True)


def _translate_online_piece(piece, target):
    if not (piece or "").strip() or not _needs_translation(piece, target): return piece
    if GoogleTranslator is None: raise RuntimeError("Optional online translator is not installed. Run: pip install deep-translator")
    translator = GoogleTranslator(source="auto", target=("ja" if target=="ja" else "en"))
    last_error=None
    for attempt in range(3):
        try: return translator.translate(piece)
        except Exception as e:
            last_error=e
            if attempt<2: time.sleep(.35*(attempt+1))
    raise RuntimeError(f"online translation service unavailable: {last_error}") from last_error


def translate_text_value(text, target, provider="online", glossary_text="", local_model=""):
    text=text or ""
    if not text.strip(): return text
    target_code="ja" if target=="ja" else "en"
    provider="online"
    local_model=""
    text=_apply_glossary_target_terms(text,target_code,glossary_text)
    translated_parts=[]
    for should_translate, segment in _protected_segments(text,target_code,glossary_text):
        if not should_translate or not _needs_translation(segment,target_code): translated_parts.append(segment); continue
        for piece in _basic_sentence_chunks(segment,target_code):
            translated_parts.append(_translate_online_piece(piece,target_code))
    result="".join(translated_parts).replace("⟦","").replace("⟧","")
    if target_code=="ja":
        result=re.sub(r"\s+([、。！？）】])",r"\1",result); result=re.sub(r"([（【])\s+",r"\1",result)
        result=to_japanese_academic_style(result)
    return result


def _collect_translatable_slots(root):
    slots=[]
    def walk(node,blocked=False):
        classes=set((node.get("class") or "").split()) if hasattr(node,"get") else set()
        here_blocked=blocked or bool({"math-inline","math-display","equation-block"}&classes) or node.get("data-latex") is not None
        tag=(node.tag or "").lower() if isinstance(node.tag,str) else ""
        if tag in {"code","pre","script","style"}: here_blocked=True
        if not here_blocked and node.text and node.text.strip(): slots.append((node,"text",node.text))
        for child in node:
            walk(child,here_blocked)
            if not here_blocked and child.tail and child.tail.strip(): slots.append((child,"tail",child.tail))
    walk(root); return slots


def translate_html_preserving_math(raw_html,target,provider="online",glossary_text="",local_model=""):
    if not (raw_html or "").strip(): return raw_html or ""
    root=lxml_html.fragment_fromstring(f"<div>{raw_html}</div>",create_parent=False)
    for node,attr,value in _collect_translatable_slots(root):
        new_text=translate_text_value(value,target,provider=provider,glossary_text=glossary_text,local_model=local_model)
        if attr=="text": node.text=new_text
        else: node.tail=new_text
    return "".join(lxml_html.tostring(child,encoding="unicode",method="html") for child in root)


def translate_paper_data(paper,target,skip_references=True,provider=None,glossary_text=None,local_model=None):
    result=deepcopy(normalize_paper_structure(paper))
    provider="online"
    local_model=""
    glossary_text=result.get("translation_glossary","") if glossary_text is None else glossary_text
    result["translation_provider"]=provider; result["translation_model"]=local_model; result["translation_glossary"]=glossary_text or ""
    if result.get("title"): result["title"]=translate_text_value(result["title"],target,provider,glossary_text,local_model)
    if result.get("abstract"): result["abstract"]=translate_text_value(result["abstract"],target,provider,glossary_text,local_model)
    if result.get("affiliations"): result["affiliations"]=translate_text_value(result["affiliations"],target,provider,glossary_text,local_model)
    ref_re=re.compile(r"^(references|bibliography|参考文献|引用文献)\s*$",re.I)
    for sec in result.get("sections",[]):
        original_title=sec.get("title",""); is_refs=bool(ref_re.match(original_title.strip()))
        if original_title and not(skip_references and is_refs): sec["title"]=translate_text_value(original_title,target,provider,glossary_text,local_model)
        if skip_references and is_refs: continue
        for blk in sec.get("blocks",[]):
            if blk.get("type") in {"image","equation"}: continue
            blk["html"]=translate_html_preserving_math(blk.get("html",""),target,provider,glossary_text,local_model)
            try: blk["plain_text"]=lxml_html.fromstring(f"<div>{blk['html']}</div>").text_content().strip()
            except Exception: blk["plain_text"]=re.sub(r"<[^>]+>","",blk.get("html",""))
        sync_section_content(sec)
    return enforce_academic_style(result) if target=="ja" else result


def is_template_metadata_table(table):
    try:
        text = " ".join((cell.text or "") for row in table.rows for cell in row.cells).lower()
    except Exception:
        return False
    cues = ["academic editor", "received:", "revised:", "accepted:", "published:", "copyright:"]
    return sum(1 for cue in cues if cue in text) >= 2


def parse_docx_to_sections(filepath, original_name, timestamp_prefix, assets, skip_elements=None, skip_indices=None):
    doc = docx.Document(filepath)
    sections = []
    current = {"id": new_id("sec"), "title": "", "blocks": []}
    image_cache = {}
    source_index = 0
    skip_elements = skip_elements or set()
    skip_indices = skip_indices or set()
    paragraph_index = -1

    for item in iter_block_items(doc):
        source_index += 1
        if isinstance(item, Paragraph):
            paragraph_index += 1
            if paragraph_index in skip_indices or item._p in skip_elements:
                continue
            if _looks_like_filename((item.text or "").strip(), original_name):
                continue
            if is_heading_paragraph(item) and (item.text or "").strip():
                if current["blocks"] or current.get("title"):
                    sync_section_content(current)
                    sections.append(current)
                current = {"id": new_id("sec"), "title": item.text.strip(), "blocks": []}
                continue
            block_type, block_html, plain_text, _, math_count = paragraph_to_html_and_assets(
                item, timestamp_prefix, image_cache, assets, original_name
            )
            if plain_text or "<img" in block_html or math_count:
                current["blocks"].append(make_block(block_type, block_html, plain_text, source_index))
        elif isinstance(item, Table):
            if source_index <= 12 and is_template_metadata_table(item):
                continue
            block_html, table_math = table_to_html(item, timestamp_prefix, image_cache, assets, original_name)
            current["blocks"].append(make_block("table", block_html, "表" + (f"（数式 {table_math}）" if table_math else ""), source_index))

    if current["blocks"] or current.get("title") or not sections:
        sync_section_content(current)
        sections.append(current)

    # Guarantee at least one editable block per section.
    for sec in sections:
        if not sec["blocks"]:
            sec["blocks"] = [make_block("paragraph", "<p></p>", "")]
        sync_section_content(sec)
    return sections


def parse_plain_text_to_sections(text):
    sections = []
    current = {"id": new_id("sec"), "title": "", "blocks": []}
    for idx, raw_line in enumerate(text.splitlines(), 1):
        line = raw_line.strip()
        if not line or line.startswith("--- Page") or line.startswith("--- Sheet"):
            continue
        if HEADING_REGEX.match(line):
            if current["blocks"]:
                sync_section_content(current)
                sections.append(current)
            current = {"id": new_id("sec"), "title": line, "blocks": []}
        else:
            current["blocks"].append(make_block("paragraph", f"<p>{html_lib.escape(line)}</p>", line, idx))
    if current["blocks"] or not sections:
        if not current["blocks"]:
            current["blocks"] = [make_block("paragraph", "<p></p>", "")]
        sync_section_content(current)
        sections.append(current)
    return sections


def safe_upload_path(src):
    if not src:
        return None
    src = src.split("?", 1)[0].split("#", 1)[0]
    if src.startswith("/uploads/"):
        filename = src[len("/uploads/"):]
    elif src.startswith("uploads/"):
        filename = src[len("uploads/"):]
    else:
        # Do not fetch remote URLs during export.
        return None
    filename = os.path.basename(filename)
    path = os.path.join(_upload_dir(), filename)
    return path if os.path.isfile(path) else None


def referenced_image_sources(paper):
    seen = set()
    sources = []
    for sec in paper.get("sections", []):
        blocks = sec.get("blocks") or split_html_into_blocks(sec.get("content", ""))
        for blk in blocks:
            raw_html = blk.get("html", "")
            try:
                frags = lxml_html.fragments_fromstring(raw_html)
                for frag in frags:
                    if isinstance(frag, str):
                        continue
                    imgs = [frag] if getattr(frag, "tag", "") == "img" else frag.xpath(".//img")
                    for img in imgs:
                        src = img.get("src")
                        if src and src not in seen:
                            seen.add(src)
                            sources.append(src)
            except Exception:
                for src in re.findall(r'<img[^>]+src=["\']([^"\']+)', raw_html, flags=re.I):
                    if src not in seen:
                        seen.add(src)
                        sources.append(src)
    return sources



# --- Project package helpers ---
PROJECT_SCHEMA_VERSION = 1
PROJECT_EXT = ".paperproj.zip"


def _clean_project_filename(value):
    value = (value or "").strip()
    value = re.sub(r"[\\/:*?\"<>|]+", "_", value)
    value = re.sub(r"\s+", " ", value).strip(" ._")
    return value[:80] or "untitled_project"


def reset_workspace(clear_uploads=True):
    """Reset the editable workspace. Project files are self-contained, so assets are project-scoped."""
    if clear_uploads:
        for entry in os.listdir(_upload_dir()):
            if entry == ".gitkeep":
                continue
            path = os.path.join(_upload_dir(), entry)
            try:
                if os.path.isdir(path):
                    shutil.rmtree(path)
                else:
                    os.remove(path)
            except Exception as e:
                print("Workspace cleanup warning:", e)
    save_assets([])
    fresh = normalize_paper_structure(json.loads(json.dumps(DEFAULT_PAPER)))
    save_paper(fresh)
    return fresh


def build_project_package(paper, ui_language="ja"):
    """Build a portable project ZIP containing manuscript JSON and all workspace assets."""
    paper = save_paper(paper or {})
    assets = load_assets()
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        project_obj = {
            "schema_version": PROJECT_SCHEMA_VERSION,
            "app": "Academic Paper Editor",
            "saved_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "ui_language": ui_language if ui_language in {"ja", "en"} else "ja",
            "paper": paper,
        }
        zf.writestr("project.json", json.dumps(project_obj, ensure_ascii=False, indent=2))
        zf.writestr("assets_meta.json", json.dumps(assets, ensure_ascii=False, indent=2))
        for asset in assets:
            filename = os.path.basename(asset.get("id", ""))
            if not filename:
                continue
            path = os.path.join(_upload_dir(), filename)
            if os.path.isfile(path):
                zf.write(path, arcname=f"assets/{filename}")
    buf.seek(0)
    return buf


def load_project_package(file_storage):
    """Load a project package after validating all archive member paths."""
    raw = file_storage.read()
    try:
        zf = zipfile.ZipFile(BytesIO(raw), "r")
    except zipfile.BadZipFile as e:
        raise ValueError("Invalid project ZIP") from e
    with zf:
        names = zf.namelist()
        if "project.json" not in names:
            raise ValueError("project.json is missing")
        for name in names:
            p = Path(name)
            if p.is_absolute() or ".." in p.parts:
                raise ValueError("Unsafe path in project file")
        obj = json.loads(zf.read("project.json").decode("utf-8"))
        if isinstance(obj, dict) and "paper" in obj:
            paper = normalize_paper_structure(obj.get("paper") or {})
            ui_language = obj.get("ui_language", "ja")
        else:
            # Forward-compatible fallback for early JSON-only project packages.
            paper = normalize_paper_structure(obj)
            ui_language = "ja"
        try:
            assets = json.loads(zf.read("assets_meta.json").decode("utf-8")) if "assets_meta.json" in names else []
        except Exception:
            assets = []

        # Keep orphaned upload files so older local snapshots do not lose their image references.
        # The active asset list itself is replaced by this project's metadata.
        reset_workspace(clear_uploads=False)
        valid_assets = []
        for asset in assets if isinstance(assets, list) else []:
            filename = os.path.basename(str(asset.get("id", "")))
            member = f"assets/{filename}"
            if filename and member in names:
                with open(os.path.join(_upload_dir(), filename), "wb") as f:
                    f.write(zf.read(member))
                asset = dict(asset)
                asset["id"] = filename
                valid_assets.append(asset)
        save_assets(valid_assets)
        paper = save_paper(paper)
        return paper, (ui_language if ui_language in {"ja", "en"} else "ja")

# --- Routing ---
def _json_payload():
    data = request.get_json(silent=True)
    return data if isinstance(data, dict) else {}


@app.errorhandler(Exception)
def handle_unexpected_error(exc):
    """Return useful API diagnostics instead of an opaque HTML 500 page."""
    if isinstance(exc, HTTPException):
        if request.path.startswith("/api/"):
            return jsonify({"error": exc.description or exc.name, "status": exc.code}), exc.code
        return exc
    error_id = uuid.uuid4().hex[:10]
    app.logger.error("Unhandled error [%s] %s %s\n%s", error_id, request.method, request.path, traceback.format_exc())
    if request.path.startswith("/api/"):
        return jsonify({
            "error": "Internal processing error",
            "error_id": error_id,
            "detail": str(exc),
        }), 500
    return f"Internal processing error. Error ID: {error_id}", 500


@app.route("/api/health", methods=["GET"])
def api_health():
    return jsonify({
        "status": "ok",
        "version": "18",
        "runtime_root": RUNTIME_ROOT,
        "writable": _is_writable_dir(RUNTIME_ROOT),
        "translation_available": GoogleTranslator is not None,
        "online_translation_available": GoogleTranslator is not None,
        "workspace_isolated": True,
    })


@app.route("/")
def index():
    return render_template("index.html")



@app.route("/api/project/new", methods=["POST"])
def api_project_new():
    # Start a clean project without deleting files referenced by older local snapshots.
    fresh = reset_workspace(clear_uploads=False)
    return jsonify({"status": "success", "paper": fresh})


@app.route("/api/project/save", methods=["POST"])
def api_project_save():
    payload = _json_payload()
    paper = payload.get("paper") or load_paper()
    ui_language = payload.get("ui_language", "ja")
    package = build_project_package(paper, ui_language)
    title = (paper or {}).get("title", "") if isinstance(paper, dict) else ""
    download_name = _clean_project_filename(title) + PROJECT_EXT
    return send_file(package, mimetype="application/zip", as_attachment=True, download_name=download_name)


@app.route("/api/project/open", methods=["POST"])
def api_project_open():
    file = request.files.get("file")
    if not file or not file.filename:
        return jsonify({"error": "No project file"}), 400
    try:
        paper, ui_language = load_project_package(file)
        return jsonify({"status": "success", "paper": paper, "ui_language": ui_language, "assets": load_assets()})
    except Exception as e:
        return jsonify({"error": f"Project load failed: {e}"}), 400

@app.route('/uploads/<path:filename>')
def serve_uploads(filename):
    return send_from_directory(_upload_dir(), filename)


@app.route("/api/paper", methods=["GET", "POST"])
def api_paper():
    if request.method == "POST":
        normalized = save_paper(_json_payload())
        return jsonify({"status": "success", "paper": normalized})
    return jsonify(load_paper())


@app.route("/api/import_document", methods=["POST"])
def import_document():
    """Import a document into ordered, paragraph-level editable blocks with inline Word images."""
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({"error": "No file"}), 400

    safe_name = secure_filename(file.filename) or "document"
    timestamp_prefix = str(int(time.time() * 1000))
    filename = f"{timestamp_prefix}_{safe_name}"
    filepath = os.path.join(_upload_dir(), filename)
    file.save(filepath)

    assets = load_assets()
    ext = Path(filename).suffix.lower()
    is_image = ext in IMAGE_EXTS
    extracted_text = ""

    register_asset(
        assets,
        filename,
        file.filename,
        "[画像ファイルです]" if is_image else "[文書ファイル]",
        is_image,
    )

    paper = load_paper()
    paper['title'] = ""
    paper['abstract'] = ""
    paper['authors'] = ""
    paper['affiliations'] = ""

    if is_image:
        block = make_block("image", f'<div class="figure-block"><img src="/uploads/{filename}" alt="{html_lib.escape(file.filename)}" class="doc-image"></div>', "")
        paper['sections'] = [{"id": new_id("sec"), "title": "", "blocks": [block]}]
        sync_section_content(paper['sections'][0])
    elif ext == ".docx":
        try:
            extracted_text = extract_text(filepath, file.filename)
            # Update source asset text preview.
            for a in assets:
                if a.get("id") == filename:
                    a["extracted_text"] = extracted_text
                    break
            parsed_doc = docx.Document(filepath)
            front = extract_docx_front_matter(parsed_doc, file.filename)
            paper['title'] = front.get('title', '')
            paper['authors'] = front.get('authors', '')
            paper['affiliations'] = front.get('affiliations', '')
            paper['abstract'] = front.get('abstract', '')
            paper['sections'] = parse_docx_to_sections(filepath, file.filename, timestamp_prefix, assets, front.get('skip_elements'), front.get('skip_indices'))
        except Exception as e:
            return jsonify({"error": f"Word解析エラー: {e}"}), 422
    else:
        extracted_text = extract_text(filepath, file.filename)
        for a in assets:
            if a.get("id") == filename:
                a["extracted_text"] = extracted_text
                break
        paper['sections'] = parse_plain_text_to_sections(extracted_text)

    save_assets(assets)
    paper = save_paper(paper)
    return jsonify({"status": "success", "paper": paper})


@app.route("/api/translation_status", methods=["GET"])
def api_translation_status():
    return jsonify({
        "status": "ok",
        "default_provider": "online",
        "online_available": GoogleTranslator is not None,
        "service": "Google Translate via deep-translator (API key not required)",
    })

@app.route("/api/translate", methods=["POST"])
def api_translate():
    payload = _json_payload()
    target = payload.get("target", "en")
    provider = "online"
    glossary_text = payload.get("glossary", "")
    local_model = ""
    if target not in {"ja", "en"}:
        return jsonify({"error": "Target language must be ja or en"}), 400
    try:
        if "paper" in payload:
            source_paper = payload.get("paper") or {}
            if not glossary_text and isinstance(source_paper, dict):
                glossary_text = source_paper.get("translation_glossary", "")
            translated = translate_paper_data(
                source_paper,
                target,
                bool(payload.get("skip_references", True)),
                provider=provider,
                glossary_text=glossary_text,
                local_model=local_model,
            )
            return jsonify({"status": "success", "paper": translated, "provider": provider})
        if "html" in payload:
            return jsonify({"status": "success", "html": translate_html_preserving_math(payload.get("html", ""), target, provider, glossary_text, local_model), "provider": provider})
        return jsonify({"status": "success", "text": translate_text_value(payload.get("text", ""), target, provider, glossary_text, local_model), "provider": provider})
    except RuntimeError as e:
        # A missing local model/package or unavailable online translator is a service-availability problem, not a server crash.
        return jsonify({"error": f"Translation unavailable: {e}", "provider": provider}), 503
    except Exception as e:
        return jsonify({"error": f"Translation failed: {e}", "provider": provider}), 502


@app.route("/api/snapshot", methods=["POST"])
def create_snapshot():
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    filename = f"snapshot_{timestamp}.json"
    payload = _json_payload()
    snapshot_obj = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "note": payload.get("note", "手動保存"),
        "data": load_paper(),
    }
    target = os.path.join(_snapshot_dir(), filename)
    tmp = target + f".tmp_{uuid.uuid4().hex}"
    with _FILE_LOCK:
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(snapshot_obj, f, ensure_ascii=False, indent=2)
            f.flush(); os.fsync(f.fileno())
        os.replace(tmp, target)
    return jsonify({"status": "success"})


@app.route("/api/snapshots", methods=["GET"])
def list_snapshots():
    snapshots = []
    for f in sorted(os.listdir(_snapshot_dir()), reverse=True):
        if f.endswith(".json"):
            try:
                with open(os.path.join(_snapshot_dir(), f), "r", encoding="utf-8") as file:
                    obj = json.load(file)
                    snapshots.append({
                        "filename": f,
                        "timestamp": obj.get("timestamp", f),
                        "note": obj.get("note", ""),
                    })
            except Exception:
                pass
    return jsonify(snapshots)


@app.route("/api/snapshot/restore", methods=["POST"])
def restore_snapshot():
    requested = os.path.basename(_json_payload().get("filename", ""))
    filepath = os.path.join(_snapshot_dir(), requested)
    if os.path.exists(filepath):
        with open(filepath, "r", encoding="utf-8") as f:
            obj = json.load(f)
            data = save_paper(obj["data"])
            return jsonify({"status": "success", "data": data})
    return jsonify({"error": "Not found"}), 404


@app.route("/api/assets", methods=["GET"])
def api_get_assets():
    return jsonify(load_assets())


@app.route("/api/upload", methods=["POST"])
def api_upload():
    file = request.files.get('file')
    if not file or file.filename == '':
        return jsonify({"error": "No file"}), 400

    safe_name = secure_filename(file.filename) or "asset"
    timestamp_prefix = str(int(time.time() * 1000))
    filename = f"{timestamp_prefix}_{safe_name}"
    filepath = os.path.join(_upload_dir(), filename)
    file.save(filepath)

    assets = load_assets()
    is_image = Path(filename).suffix.lower() in IMAGE_EXTS
    extracted_text = "[画像ファイルです]" if is_image else extract_text(filepath, file.filename)
    asset = register_asset(assets, filename, file.filename, extracted_text, is_image)
    save_assets(assets)
    return jsonify({"status": "success", "asset": asset})


@app.route("/api/import_material", methods=["POST"])
def import_material():
    """Parse an uploaded material using the same document parser as the initial import.

    The current manuscript metadata is not replaced. Parsed sections/blocks are returned to the
    client for insertion at the current position, while embedded images and the source file are
    registered in the asset library.
    """
    file = request.files.get("file")
    if not file or file.filename == "":
        return jsonify({"error": "No file"}), 400
    safe_name = secure_filename(file.filename) or "document"
    timestamp_prefix = str(int(time.time() * 1000))
    filename = f"{timestamp_prefix}_{safe_name}"
    filepath = os.path.join(_upload_dir(), filename)
    file.save(filepath)
    assets = load_assets()
    ext = Path(filename).suffix.lower()
    is_image = ext in IMAGE_EXTS
    source_asset = register_asset(assets, filename, file.filename, "[画像ファイルです]" if is_image else "[文書ファイル]", is_image)
    metadata = {"title":"", "authors":"", "affiliations":"", "abstract":""}
    try:
        if is_image:
            sections = [{"id": new_id("sec"), "title":"", "blocks":[make_block("image", f'<div class="figure-block"><img src="/uploads/{filename}" alt="{html_lib.escape(file.filename)}" class="doc-image"></div>', "")]}]
            sync_section_content(sections[0])
        elif ext == ".docx":
            extracted_text = extract_text(filepath, file.filename)
            source_asset["extracted_text"] = extracted_text
            parsed_doc = docx.Document(filepath)
            front = extract_docx_front_matter(parsed_doc, file.filename)
            metadata = {k: front.get(k, "") for k in ("title","authors","affiliations","abstract")}
            sections = parse_docx_to_sections(filepath, file.filename, timestamp_prefix, assets, front.get("skip_elements"), front.get("skip_indices"))
        else:
            extracted_text = extract_text(filepath, file.filename)
            source_asset["extracted_text"] = extracted_text
            sections = parse_plain_text_to_sections(extracted_text)
        save_assets(assets)
        block_count = sum(len(sec.get("blocks",[])) for sec in sections)
        math_count = sum((blk.get("html","").count("data-omml-b64=")) for sec in sections for blk in sec.get("blocks",[]))
        image_count = sum((blk.get("html","").count("<img")) for sec in sections for blk in sec.get("blocks",[]))
        return jsonify({"status":"success", "sections":sections, "metadata":metadata, "asset":source_asset,
                        "summary":{"sections":len(sections),"blocks":block_count,"equations":math_count,"images":image_count}})
    except Exception as e:
        return jsonify({"error": f"Material parse error: {e}"}), 422


@app.route("/api/assets/<asset_id>/parse", methods=["POST"])
def api_parse_existing_asset(asset_id):
    """Re-parse a stored asset using the same structure-aware parser used by initial import."""
    asset_id = os.path.basename(asset_id)
    filepath = os.path.join(_upload_dir(), asset_id)
    if not os.path.isfile(filepath):
        return jsonify({"error":"Not found"}), 404
    asset = next((a for a in load_assets() if a.get("id") == asset_id), None)
    if not asset:
        return jsonify({"error":"Asset metadata not found"}), 404
    original_name = asset.get("original_name") or asset_id
    timestamp_prefix = str(int(time.time() * 1000))
    assets = load_assets()
    ext = Path(filepath).suffix.lower()
    try:
        if ext in IMAGE_EXTS:
            sections=[{"id":new_id("sec"),"title":"","blocks":[make_block("image", f'<div class="figure-block"><img src="/uploads/{asset_id}" alt="{html_lib.escape(original_name)}" class="doc-image"></div>',"")]}]
            sync_section_content(sections[0])
        elif ext == ".docx":
            parsed_doc = docx.Document(filepath)
            front = extract_docx_front_matter(parsed_doc, original_name)
            sections = parse_docx_to_sections(filepath, original_name, timestamp_prefix, assets, front.get("skip_elements"), front.get("skip_indices"))
        else:
            sections = parse_plain_text_to_sections(extract_text(filepath, original_name))
        save_assets(assets)
        block_count=sum(len(sec.get("blocks",[])) for sec in sections)
        math_count=sum(blk.get("html","").count("data-omml-b64=") for sec in sections for blk in sec.get("blocks",[]))
        image_count=sum(blk.get("html","").count("<img") for sec in sections for blk in sec.get("blocks",[]))
        return jsonify({"status":"success","sections":sections,"summary":{"sections":len(sections),"blocks":block_count,"equations":math_count,"images":image_count}})
    except Exception as e:
        return jsonify({"error":f"Asset parse error: {e}"}), 422


@app.route("/api/assets/<asset_id>", methods=["DELETE"])
def api_delete_asset(asset_id):
    asset_id = os.path.basename(asset_id)
    save_assets([a for a in load_assets() if a.get('id') != asset_id])
    filepath = os.path.join(_upload_dir(), asset_id)
    if os.path.exists(filepath):
        os.remove(filepath)
    return jsonify({"status": "success"})


@app.route("/api/assets/<asset_id>/download", methods=["GET"])
def api_download_asset(asset_id):
    asset_id = os.path.basename(asset_id)
    filepath = os.path.join(_upload_dir(), asset_id)
    if not os.path.isfile(filepath):
        return jsonify({"error": "Not found"}), 404
    original = next((a.get("original_name") for a in load_assets() if a.get("id") == asset_id), asset_id)
    return send_file(filepath, as_attachment=True, download_name=secure_filename(original) or asset_id)


# --- Export helpers ---
def decode_omml_from_node(node):
    b64 = node.get("data-omml-b64") if hasattr(node, "get") else None
    if not b64:
        return None
    try:
        xml = base64.urlsafe_b64decode(b64.encode("ascii"))
        return etree.fromstring(xml)
    except Exception as e:
        print("OMML decode error:", e)
        return None


def append_math_node_to_docx(paragraph, node):
    omml = decode_omml_from_node(node)
    if omml is not None:
        try:
            # Inline OMML is valid inside a Word paragraph. For oMathPara, append its oMath children.
            if _local(omml) == "oMathPara":
                math_children = omml.findall(f'{{{MATH_NS}}}oMath')
                if math_children:
                    for child in math_children:
                        paragraph._p.append(deepcopy(child))
                else:
                    paragraph._p.append(deepcopy(omml))
            else:
                paragraph._p.append(deepcopy(omml))
            return True
        except Exception as e:
            print("OMML append error:", e)
    latex = node.get("data-latex") or node.text_content().strip()
    if latex:
        run = paragraph.add_run(latex)
        run.font.name = "Cambria Math"
        return True
    return False


def add_inline_node_to_docx(paragraph, node, preset_key="generic"):
    if isinstance(node, str):
        if node:
            paragraph.add_run(node)
        return
    tag = (node.tag or "").lower() if isinstance(node.tag, str) else ""
    classes = set((node.get("class") or "").split()) if hasattr(node, "get") else set()

    if tag == "span" and ({"math-inline", "math-display"} & classes or node.get("data-latex") is not None):
        append_math_node_to_docx(paragraph, node)
        return
    if tag == "img":
        path = safe_upload_path(node.get("src"))
        if path:
            try:
                run = paragraph.add_run()
                run.add_picture(path, width=Inches(3.15 if preset_key == "ieee" else 5.2))
            except Exception as e:
                paragraph.add_run(f"[画像エラー: {e}]")
        else:
            paragraph.add_run("[画像が見つかりません]")
        return
    if tag == "br":
        paragraph.add_run().add_break()
        return

    if node.text:
        run = paragraph.add_run(node.text)
        if tag in {"strong", "b"}:
            run.bold = True
        if tag in {"em", "i"}:
            run.italic = True
    for child in node:
        add_inline_node_to_docx(paragraph, child, preset_key)
        if child.tail:
            paragraph.add_run(child.tail)


def _clear_docx_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if _local(child) != 'pPr':
            p.remove(child)


def _fill_cell_from_html(cell, cell_node, preset_key="generic"):
    paragraphs = cell_node.xpath('./p|./div|./figure')
    if not paragraphs:
        paragraphs = [cell_node]
    first = True
    for frag in paragraphs:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        _clear_docx_paragraph(p)
        if frag.text:
            p.add_run(frag.text)
        for child in frag:
            add_inline_node_to_docx(p, child, preset_key)
            if child.tail:
                p.add_run(child.tail)


def _set_table_borders_none(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "nil")


def append_html_block_to_docx(doc, raw_html, preset_key="generic"):
    if not raw_html or not raw_html.strip():
        doc.add_paragraph("")
        return
    try:
        fragments = lxml_html.fragments_fromstring(raw_html)
    except Exception:
        doc.add_paragraph(re.sub(r"<[^>]+>", "", raw_html))
        return

    for frag in fragments:
        if isinstance(frag, str):
            if frag.strip():
                doc.add_paragraph(frag)
            continue
        tag = (frag.tag or "").lower() if isinstance(frag.tag, str) else ""
        classes = set((frag.get("class") or "").split()) if hasattr(frag, "get") else set()
        if tag == "table":
            rows = frag.xpath(".//tr")
            max_cols = max([len(r.xpath("./th|./td")) for r in rows] or [1])
            table = doc.add_table(rows=0, cols=max_cols)
            if "equation-table" in classes:
                _set_table_borders_none(table)
            else:
                table.style = "Table Grid"
            table.autofit = preset_key != "ieee"
            for row_node in rows:
                cells_nodes = row_node.xpath("./th|./td")
                cells = table.add_row().cells
                for i, cell_node in enumerate(cells_nodes[:max_cols]):
                    _fill_cell_from_html(cells[i], cell_node, preset_key)
            if preset_key == "ieee":
                col_width = Inches(3.15 / max_cols)
                for col in table.columns:
                    col.width = col_width
                for row in table.rows:
                    for cell in row.cells:
                        cell.width = col_width
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(6.8)
        elif tag in {"ul", "ol"}:
            style = "List Bullet" if tag == "ul" else "List Number"
            for li in frag.xpath("./li"):
                p = doc.add_paragraph(style=style)
                add_inline_node_to_docx(p, li, preset_key)
        elif tag == "img":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_node_to_docx(p, frag, preset_key)
        elif tag in {"div", "figure"} and frag.xpath(".//img") and not frag.text_content().strip():
            for img in frag.xpath(".//img"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_node_to_docx(p, img, preset_key)
        else:
            p = doc.add_paragraph()
            if "equation-block" in classes:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if frag.text:
                p.add_run(frag.text)
            for child in frag:
                add_inline_node_to_docx(p, child, preset_key)
                if child.tail:
                    p.add_run(child.tail)


def set_section_columns(section, num=1, space_twips=360):
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    cols_el = cols[0] if cols else OxmlElement('w:cols')
    if not cols:
        sectPr.append(cols_el)
    cols_el.set(qn('w:num'), str(num))
    if num > 1:
        cols_el.set(qn('w:space'), str(space_twips))


def apply_publisher_docx_style(doc, preset_key, content_section=None):
    preset = PUBLISHER_PRESETS.get(preset_key, PUBLISHER_PRESETS['generic'])
    for sec in doc.sections:
        if preset['page'] == 'Letter':
            sec.page_width = Inches(8.5); sec.page_height = Inches(11)
        else:
            sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
        top, right, bottom, left = preset['margins_cm']
        sec.top_margin = Cm(top); sec.right_margin = Cm(right); sec.bottom_margin = Cm(bottom); sec.left_margin = Cm(left)
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = preset['font']
    normal.font.size = Pt(preset['font_size'])
    normal.paragraph_format.line_spacing = preset['line_spacing']
    normal.paragraph_format.space_after = Pt(3 if preset_key == 'ieee' else 6)
    for name, size in [('Title', 18 if preset_key == 'ieee' else 16), ('Heading 1', 11 if preset_key == 'ieee' else 13), ('Heading 2', 10 if preset_key == 'ieee' else 12)]:
        if name in styles:
            styles[name].font.name = preset['font']
            styles[name].font.size = Pt(size)
    if content_section is not None:
        set_section_columns(content_section, preset['columns'], 360 if preset_key == 'ieee' else 480)


def latex_escape_text(text):
    if text is None:
        return ''
    repl = {'\\': r'\textbackslash{}', '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_', '{': r'\{', '}': r'\}', '~': r'\textasciitilde{}', '^': r'\textasciicircum{}'}
    return ''.join(repl.get(ch, ch) for ch in str(text))


def html_node_to_latex(node):
    if isinstance(node, str):
        return latex_escape_text(node)
    tag = (node.tag or '').lower() if isinstance(node.tag, str) else ''
    classes = set((node.get('class') or '').split()) if hasattr(node, 'get') else set()
    if tag == 'span' and ({'math-inline', 'math-display'} & classes or node.get('data-latex') is not None):
        latex = html_lib.unescape(node.get('data-latex') or node.text_content().strip())
        display = node.get('data-math-display') == '1' or 'math-display' in classes
        return f"\\[{latex}\\]" if display else f"${latex}$"
    if tag == 'img':
        src = node.get('src') or ''
        fn = os.path.basename(src)
        alt = latex_escape_text(node.get('alt') or '')
        return f"\\begin{{figure}}[htbp]\n\\centering\n\\includegraphics[width=0.85\\linewidth]{{images/{fn}}}\n\\caption{{{alt}}}\n\\end{{figure}}\n"
    if tag == 'br':
        return '\\\\\n'
    if tag == 'table':
        rows = node.xpath('.//tr')
        ncols = max([len(r.xpath('./th|./td')) for r in rows] or [1])
        out = [f"\\begin{{tabular}}{{{'l'*ncols}}}", r"\hline"]
        for row in rows:
            vals = []
            for cell in row.xpath('./th|./td'):
                vals.append(''.join(html_node_to_latex(x) for x in cell.xpath('./node()')))
            vals += [''] * (ncols - len(vals))
            out.append(' & '.join(vals) + r' \\')
        out += [r'\hline', r'\end{tabular}']
        return '\n'.join(out) + '\n'
    prefix = suffix = ''
    if tag in {'strong', 'b'}: prefix, suffix = r'\textbf{', '}'
    elif tag in {'em', 'i'}: prefix, suffix = r'\textit{', '}'
    parts = []
    if node.text:
        parts.append(latex_escape_text(node.text))
    for child in node:
        parts.append(html_node_to_latex(child))
        if child.tail:
            parts.append(latex_escape_text(child.tail))
    body = ''.join(parts)
    if tag in {'p', 'div', 'figure', 'li'}:
        body += '\n\n'
    return prefix + body + suffix


def html_block_to_latex(raw_html):
    try:
        frags = lxml_html.fragments_fromstring(raw_html or '')
    except Exception:
        return latex_escape_text(re.sub(r'<[^>]+>', '', raw_html or ''))
    return ''.join(html_node_to_latex(f) for f in frags)


def _block_fallback_text(block):
    text = (block or {}).get("plain_text", "")
    if text:
        return text
    return re.sub(r"<[^>]+>", "", (block or {}).get("html", "")).strip()


def latex_preamble_for_preset(paper):
    key = paper.get('format_preset', 'generic')
    p = PUBLISHER_PRESETS.get(key, PUBLISHER_PRESETS['generic'])
    cls, opts = p['latex_class'], p['latex_options']
    lines = [f"\\documentclass[{opts}]{{{cls}}}", "\\usepackage{graphicx}", "\\usepackage{amsmath,amssymb}", "\\usepackage{booktabs}"]
    if key == 'springer_nature':
        lines.append('% Springer Nature: sn-jnl.cls is supplied by the official journal template package.')
    elif key in {'wiley', 'taylor_francis', 'mdpi'}:
        lines.append('% General publisher preset. Replace with the target journal-specific class/template when required.')
    lines += [f"\\title{{{latex_escape_text(paper.get('title',''))}}}", f"\\author{{{latex_escape_text(paper.get('authors',''))}}}", "\\date{}", "", "\\begin{document}", "\\maketitle"]
    if paper.get('abstract'):
        lines += ["\\begin{abstract}", latex_escape_text(paper.get('abstract','')), "\\end{abstract}", ""]
    return '\n'.join(lines) + '\n'

@app.route("/api/export/<format_type>", methods=["GET"])
def export_paper(format_type):
    paper = load_paper()
    preset_key = paper.get('format_preset', 'generic')
    preset = PUBLISHER_PRESETS.get(preset_key, PUBLISHER_PRESETS['generic'])

    if format_type == "latex":
        tex = latex_preamble_for_preset(paper)
        for sec in paper.get("sections", []):
            if sec.get('title'):
                tex += f"\\section{{{latex_escape_text(sec.get('title',''))}}}\n"
            for blk in sec.get("blocks", []):
                try:
                    tex += html_block_to_latex(blk.get("html", "")) + "\n"
                except Exception as e:
                    app.logger.warning("LaTeX block fallback: %s", e)
                    tex += latex_escape_text(_block_fallback_text(blk)) + "\n\n"
        if paper.get('bib_data'):
            tex += "\n% BibTeX data stored in the editor:\n% " + paper.get('bib_data','').replace('\n', '\n% ') + "\n"
        tex += "\\end{document}\n"
        return Response(tex, mimetype="text/plain; charset=utf-8", headers={"Content-Disposition": f"attachment;filename=paper_{preset_key}.tex"})

    if format_type == "latex_zip":
        tex_response = export_paper("latex")
        tex_bytes = tex_response.get_data()
        out = BytesIO()
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"paper_{preset_key}.tex", tex_bytes)
            for src in referenced_image_sources(paper):
                path = safe_upload_path(src)
                if path:
                    zf.write(path, arcname=f"images/{os.path.basename(path)}")
            zf.writestr("FORMAT_NOTE.txt", (
                f"Preset: {preset['label']}\n"
                "This is a publisher-level authoring preset. Final submission requirements may differ by journal. "
                "Use the target journal's official Instructions for Authors/template as the final authority.\n"
            ))
        out.seek(0)
        return send_file(out, as_attachment=True, download_name=f"paper_{preset_key}_latex.zip", mimetype="application/zip")

    if format_type == "docx":
        doc = docx.Document()
        # Front matter stays single-column. IEEE body switches to two columns after abstract.
        title_p = doc.add_paragraph(style='Title')
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.add_run(paper.get('title', 'Untitled'))
        if paper.get('authors'):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(paper.get('authors', ''))
        if paper.get('affiliations'):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(paper.get('affiliations', '')); r.italic = True
        if paper.get('abstract'):
            if preset_key == 'ieee':
                p = doc.add_paragraph(); p.add_run('Abstract—').bold = True; p.add_run(paper.get('abstract', ''))
            else:
                doc.add_heading('Abstract', level=1)
                doc.add_paragraph(paper.get('abstract', ''))

        content_section = doc.sections[0]
        if preset.get('columns', 1) > 1:
            content_section = doc.add_section(WD_SECTION.CONTINUOUS)
        apply_publisher_docx_style(doc, preset_key, content_section)

        for sec in paper.get("sections", []):
            if sec.get('title'):
                doc.add_heading(sec.get('title', ''), level=1)
            for blk in sec.get("blocks", []):
                try:
                    append_html_block_to_docx(doc, blk.get("html", ""), preset_key)
                except Exception as e:
                    app.logger.warning("DOCX block fallback: %s", e)
                    fallback = _block_fallback_text(blk)
                    if fallback:
                        doc.add_paragraph(fallback)

        f = BytesIO()
        doc.save(f)
        f.seek(0)
        return send_file(
            f,
            as_attachment=True,
            download_name=f"paper_{preset_key}_with_math_images.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if format_type == "images_zip":
        out = BytesIO()
        added = set()
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
            for src in referenced_image_sources(paper):
                path = safe_upload_path(src)
                if path and path not in added:
                    zf.write(path, arcname=os.path.basename(path))
                    added.add(path)
        out.seek(0)
        return send_file(out, as_attachment=True, download_name="paper_images.zip", mimetype="application/zip")

    if format_type in ["pdf_print", "html"]:
        return render_template("export_template.html", paper=paper, auto_print=(format_type == "pdf_print"), preset=preset_key)

    return jsonify({"error": "Invalid format"}), 400


if __name__ == "__main__":
    app.run(
        host="0.0.0.0",
        port=int(os.environ.get("PORT", "5000")),
        debug=os.environ.get("FLASK_DEBUG", "0") == "1",
    )
